import fitz  # import for the text extraction for pdf
from transformers import BertTokenizer, BertModel
import torch
import os
import numpy as np
from sentence_transformers import SentenceTransformer
import re
from pathlib import Path
from annoy import AnnoyIndex
import random

from docx import Document

import pandas as pd
import pyarrow as pa
import pyarrow.parquet as pq

from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs

df = None
annoy_index = None


def get_text():
    '''full_text = ""
    for i in range(1,46): 
        document_name = "/Users/mariam/Desktop/nis2p/word/" + str(i) + ".docx"
        document = Document(document_name)
        for paragraph in document.paragraphs:
            full_text += paragraph.text + "\n"  
    return full_text.strip()'''
    
    full_text = ""
    document = Document("try.docx")
    for paragraph in document.paragraphs:
        full_text += paragraph.text + "\n"  
    return full_text.strip()

def chunk_text(text):
    sentences = re.split(r"(?<=[.؟!?،؛:])\s+", text)
    chunks = []
    for sentence in sentences:
        sentence = sentence.strip()
        if sentence:
            chunks.append(sentence)
    return chunks

     
def generate_embeddings(text_chunks):
    model = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")
    embeddings = []
    print("hello")
    for chunk in text_chunks:
        emb = list(model.encode(chunk))
        embeddings.append(emb)
    return embeddings


'''def do_work():
    # Interactive user input and query processing
    global df, annoy_index
    user_query = input("Please enter a question: ")
    query_chunks = chunk_text(user_query)
    query_embedding = generate_embeddings(query_chunks)
    query_embedding = np.array(query_embedding)
    query_embedding = query_embedding.squeeze()

    near_index = annoy_index.get_nns_by_vector(query_embedding, 1)[0]
    max_text = df.loc[near_index, 'text']
    print("\nAnswer: ", max_text)'''


# Initialize df and annoy_index globally before server starts
file_path = "/Users/mariam/Desktop/nis2p/output_embeddings.parquet"
if os.path.exists(file_path): 
    df = pd.read_parquet(file_path)
else: 
    full_text = get_text()
    text_chunks = chunk_text(full_text)
    embeddings = generate_embeddings(text_chunks)
    df = pd.DataFrame({'text': text_chunks, 'embedding': embeddings})
    df.to_parquet(file_path)

f = len(np.array(df.loc[0, 'embedding']))
annoy_index = AnnoyIndex(f, 'dot')
for i, row in df.iterrows():
    row_vector = np.array(row['embedding'])
    annoy_index.add_item(i, row_vector)
annoy_index.build(20)


class query_server(BaseHTTPRequestHandler): #https://www.youtube.com/watch?v=DeFST8tvtuI&t=461s
    def do_GET(self):
        parsed = urlparse(self.path)
        query = parse_qs(parsed.query).get('q', [''])[0]
        if not query:
            self.send_response(400)
            self.end_headers()
            return

        query_embedding = generate_embeddings([query])[0]
        index = annoy_index.get_nns_by_vector(query_embedding, 1)[0]
        result = df.loc[index, 'text']

        self.send_response(200)
        self.send_header('Content-Type', 'text/plain')
        self.end_headers()
        self.wfile.write(result.encode('utf-8'))


def run_server(port=8000):
    server = HTTPServer(('', port), query_server)
    print(f"http://localhost:{port}")
    server.serve_forever()

#do_work()
run_server()
